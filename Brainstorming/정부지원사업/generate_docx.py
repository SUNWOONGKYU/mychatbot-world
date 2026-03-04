# -*- coding: utf-8 -*-
"""
재도전성공패키지 사업계획서 - MD 파일 충실 반영 DOCX 생성 (v3)
원본: 사업계획서_재도전성공패키지_MyChabbotWorld.md
v3 변경: page break, bold in cells, 아키텍처 고정폭, O/X 중앙정렬, 누락 문구 추가
"""
import sys
sys.stdout.reconfigure(encoding='utf-8')
import re
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

doc = Document()

# ── 페이지 설정 ──
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# ── 스타일 설정 ──
style = doc.styles['Normal']
style.font.name = 'Malgun Gothic'
style.font.size = Pt(10)
style.paragraph_format.space_after = Pt(2)
style.paragraph_format.line_spacing = 1.15
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')

def set_font(run, name='Malgun Gothic'):
    run.font.name = name
    run.element.rPr.rFonts.set(qn('w:eastAsia'), name)

def add_heading_styled(text, level=1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = True
    if level == 0:
        run.font.size = Pt(14)
    elif level == 1:
        run.font.size = Pt(12)
    elif level == 2:
        run.font.size = Pt(11)
    else:
        run.font.size = Pt(10)
    set_font(run)
    return p

def add_body(text, bold=False, size=10, indent=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    set_font(run)
    return p

def add_monospace(text, size=8):
    """고정폭 폰트로 출력 (아키텍처 다이어그램용)"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(text)
    run.font.size = Pt(size)
    set_font(run, 'Consolas')
    return p

def add_horizontal_rule():
    """수평선 추가"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._element.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        '  <w:bottom w:val="single" w:sz="6" w:space="1" w:color="999999"/>'
        '</w:pBdr>'
    )
    pPr.append(pBdr)

def _write_rich_cell(cell, text, font_size=9, alignment=None):
    """표 셀에 **bold** 마크업을 파싱하여 rich text로 작성"""
    cell.text = ''
    p = cell.paragraphs[0]
    if alignment:
        p.alignment = alignment
    # **text** 패턴 분할
    parts = re.split(r'(\*\*.*?\*\*)', str(text))
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run.bold = True
        else:
            run = p.add_run(part)
        run.font.size = Pt(font_size)
        set_font(run)

def add_table(headers, rows, col_widths=None, col_aligns=None):
    """
    표 추가. col_aligns: 컬럼별 정렬 리스트 ('left','center','right')
    셀 텍스트에 **bold** 마크업 지원.
    """
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    align_map = {
        'left': WD_ALIGN_PARAGRAPH.LEFT,
        'center': WD_ALIGN_PARAGRAPH.CENTER,
        'right': WD_ALIGN_PARAGRAPH.RIGHT,
    }

    # 헤더
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        set_font(run)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="D9E2F3"/>')
        cell._element.get_or_add_tcPr().append(shading)

    # 데이터
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            a = None
            if col_aligns and c_idx < len(col_aligns) and col_aligns[c_idx]:
                a = align_map.get(col_aligns[c_idx])
            _write_rich_cell(cell, val, font_size=9, alignment=a)

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    doc.add_paragraph()
    return table

def add_bullet(text, level=0):
    prefix = "  " * level + "• "
    add_body(prefix + text, size=9, indent=True)


# ═══════════════════════════════════════════════════════════════
# 사업계획서 본문 시작 — MD 파일 충실 반영 (v3)
# ═══════════════════════════════════════════════════════════════

# ── 표지 ──
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('2026년 재도전성공패키지\n(예비)재창업기업 사업계획서')
run.bold = True
run.font.size = Pt(16)
set_font(run)
doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# □ 과제 개요
# ═══════════════════════════════════════════════════════════════
add_heading_styled('□ 과제 개요', 1)

add_table(
    ['항목', '내용'],
    [
        ['과제명 (사업명)', 'My Chatbot World — AI 챗봇의 생성·학습·성장·수익활동·피상속으로 이어지는 라이프사이클 지원 플랫폼'],
        ['기업명', '(기재 필요) ※ 예비재창업자의 경우 "예비재창업자"로 기재'],
        ['아이템(서비스)\n개요',
         '멀티 페르소나(아바타형 대외용 + 도우미형 대내용) 이원 구조와 감성 수준(1~100) 슬라이더 기반 '
         '멀티 AI 자동 라우팅을 통해, AI 챗봇의 생성→학습→성장→수익활동→피상속으로 이어지는 '
         '라이프사이클 전체를 원스톱으로 지원하는 세계 최초의 개인 맞춤형 AI 챗봇 플랫폼. '
         '(특허 출원 완료, 7대 핵심 구성요소, 청구항 22개)'],
    ],
    col_widths=[3.5, 13]
)

# ═══════════════════════════════════════════════════════════════
# □ 폐업 이력
# ═══════════════════════════════════════════════════════════════
add_heading_styled('□ 폐업 이력 (총 폐업 횟수 : ○회)', 1)

add_table(
    ['항목', '내용'],
    [
        ['기업명', '○○○○'],
        ['기업 구분', '개인 / 법인'],
        ['사업기간', '20○○.○○.○○ ~ 20○○.○○.○○'],
        ['업종(업태/종목)', '(기재 필요)'],
        ['매출 규모', '(기재 필요)'],
        ['고용 규모', '(기재 필요)'],
        ['폐업 사유', '(기재 필요)'],
    ],
    col_widths=[3.5, 13]
)


# ═══════════════════════════════════════════════════════════════
# 1. 문제인식                                    ★ PAGE BREAK ★
# ═══════════════════════════════════════════════════════════════
doc.add_page_break()
add_heading_styled('1. 문제인식 (폐업 원인 분석 및 재창업 아이템 추진 배경)', 0)

# ── 1-1. 이전 사업 폐업 원인 분석 ──
add_heading_styled('1-1. 이전 사업 폐업 원인 분석', 2)

add_body('이전 사업에서 겪은 실패를 네 가지 핵심 원인으로 분석하고, 각각에 대한 구체적 개선 방향을 도출하였습니다.', size=10)

add_table(
    ['구분', '실패 원인', '교훈 및 개선 방향'],
    [
        ['**시장 검증 부재**', '고객 니즈 확인 없이 기술 중심 개발 착수, 출시 후 시장 반응 저조',
         'MVP로 빠른 시장 검증 후 확장하는 린 스타트업 방식 채택'],
        ['**수익모델 불명확**', '명확한 과금 체계 없이 무료 운영, 지속가능한 수익 창출 실패',
         '크레딧 사용료 + 스킬 마켓 수수료 + 수익활동 중개 수수료의 다중 수익 구조 설계 완료'],
        ['**비용 통제 실패**', '자체 서버 구축 등 초기 고정비 과다 투입',
         '서버리스(Vercel) + BaaS(Supabase)로 초기 인프라 비용 월 $10 이하 달성'],
        ['**단독 운영 한계**', '기획·개발·마케팅 1인 담당으로 속도·품질 모두 저하',
         'AI 개발 도구(Claude Code, Copilot)로 1인 생산성 10배 향상 + 외부 전문 인력 협업 체계'],
    ],
    col_widths=[2.5, 6.5, 7.5]
)

add_body('핵심 교훈: "기술이 아무리 좋아도, 고객이 원하는 것을 만들지 않으면 실패한다." 이 뼈아픈 경험이 이번 재창업의 출발점입니다.', bold=True, size=10)

# ── 1-2. 재창업 아이템 추진 배경 및 필요성 ──
add_heading_styled('1-2. 재창업 아이템 추진 배경 및 필요성', 2)

add_body('(1) 시장의 구조적 문제 — 개인·소상공인 AI 챗봇 시장의 부재', bold=True, size=10)
add_body('글로벌 AI 챗봇 시장은 2023년 51억 달러에서 2028년 155억 달러(CAGR 25%)로 급성장 중이나, **기업용(B2B)이 80% 이상**을 차지하며, 개인·소상공인을 위한 시장은 사실상 **블루오션**입니다.', size=9)

add_table(
    ['구분', '기존 시장', 'My Chatbot World'],
    [
        ['**타겟**', '대기업·중견기업', '개인·소상공인·전문가'],
        ['**가격**', '월 $74~$2,500', '사용량 기반 크레딧 과금 (구독료 없음)'],
        ['**생성 소요시간**', '30분~1시간 (폼 30개 입력)', '**음성 5분** (AI가 질문, 사용자가 답변)'],
        ['**기술 진입장벽**', '높음 (코딩·설정 필요)', '**제로** (말만 하면 완성)'],
        ['**챗봇 라이프사이클**', '생성만 지원', '**생성 → 학습 → 성장 → 수익활동 → 피상속** 전체 지원'],
        ['**수익활동 중개**', '없음', '**챗봇이 직접 상담·학습·업무대행 수행**, 플랫폼이 중개 (플랫폼 20% / 운영자 80%)'],
        ['**디지털 유산 관리**', '없음', '**피상속 시스템** (세계 최초)'],
    ],
    col_widths=[3.5, 5.5, 7.5]
)

add_body('(2) 기존 기술의 5가지 한계와 본 발명의 해결', bold=True, size=10)
add_body('본 아이템은 **특허 출원 완료**된 핵심 기술에 기반합니다.', size=9)
add_body('발명의 명칭: "멀티 페르소나와 감성·비용 조건 기반 멀티 AI 자동 라우팅을 구비한 AI 챗봇의 생성·학습·성장·수익활동·피상속으로 이어지는 라이프사이클 지원 플랫폼 및 방법"', size=9, indent=True)

add_table(
    ['#', '기존 기술의 한계', '본 발명의 해결'],
    [
        ['1', '기존 플랫폼(Botpress, GPTs 등)은 **단일 AI 엔진에 의존**하며, 감성적 요구에 따라 AI를 자동 선택·라우팅하는 기능 부재',
         '**감성 수준(1~100) 슬라이더 × 페르소나 조합**에 따라 복수의 AI 엔진 중 최적 AI를 자동 라우팅'],
        ['2', 'Character.ai 등은 가상 캐릭터 생성에 불과, 사용자 본인의 다면적 정체성을 **대외용과 대내용으로 체계적 분류·관리하는 구조 부재**',
         '**아바타형 페르소나(대외용 분신)**와 **도우미형 페르소나(대내용 업무·생활 도우미)**의 이원 관리 구조'],
        ['3', '기존 플랫폼은 챗봇 **생성만** 지원, 생성 이후 학습·성장·스킬 장착 등 **전체 라이프사이클 통합 관리 부재**',
         'DB축적 + 인간/AI 멘토링 + 스쿨 시스템 + 챗봇 커뮤니티 + 스킬 마켓의 **다층적 학습·성장 체계**'],
        ['4', 'MS 특허(US10,853,717B2)는 고인의 소셜 데이터를 제3자가 사후에 수집하는 기술, **본인 사전 동의 기반 피상속 메커니즘 부재**',
         '**본인이 생전에** 페르소나별·데이터별 피상속 허용/불허를 사전 설정, 비동의 항목 즉시 자동 완전 삭제'],
        ['5', 'Meta의 디지털 사후 아바타 특허도 SNS 활동 시뮬레이션에 불과, **챗봇 소유권 이전 및 피상속 후 지속 성장 개념 부재**',
         '사망 확인 시 동의 항목은 상속인에게 **소유권 이전**, 피상속 이후에도 **상속인 하에서 챗봇이 계속 학습·성장**'],
    ],
    col_widths=[0.8, 7, 8.7]
)

add_body('(3) 핵심 차별화 기술 — 특허 출원 7대 구성요소', bold=True, size=10)

add_table(
    ['#', '구성요소', '선행기술', '상세'],
    [
        ['1', '**멀티 페르소나 관리부**', '없음', '아바타형(대외용): CPA 아바타, AI전문가 아바타 등 복수 설정\n도우미형(대내용): 업무도우미 + 생활도우미'],
        ['2', '**감성 기반 멀티 AI\n라우팅부**', '없음', '감성 수준 1~100 슬라이더를 수시 조정\n→ 팩트 중심(지적 응답) ~ 공감·위로(감성적 응답) 자동 전환'],
        ['3', '**챗봇 자동 생성부**', '없음', 'AI가 음성으로 질문하고 사용자가 음성으로 답하는 인터뷰 방식,\n5분 이내 분신 챗봇 자동 생성'],
        ['4', '**챗봇 학습·성장 관리부**', '없음', 'DB축적 + 인간/AI 멘토링 + 커리큘럼 기반 스쿨\n+ 챗봇 커뮤니티(토론·협업·SNS) + 스킬 마켓'],
        ['5', '**피상속 관리부**', '없음', '생전 피상속 사전 설정 → 상속인 지정 → 사망 확인 트리거\n→ 동의 항목 소유권 이전 + 비동의 항목 즉시 완전 삭제\n→ 상속인 하 지속 성장'],
        ['6', '**배포부**', '—', 'URL 임베딩 + QR코드로 웹/모바일/메신저 등 모든 채널 배포'],
        ['7', '**수익활동 중개부**', '없음', '챗봇이 상담·학습·업무대행 등 수익활동을 능동적으로 수행\n→ 플랫폼 수수료 20% 자동 수취 + 운영자 80% 정산\n(예: CPA 챗봇 세무 상담, 마케팅 챗봇 SNS 업무대행)'],
    ],
    col_widths=[0.7, 3.3, 1.5, 11]
)


# ═══════════════════════════════════════════════════════════════
# 2. 실현가능성                                  ★ PAGE BREAK ★
# ═══════════════════════════════════════════════════════════════
doc.add_page_break()
add_heading_styled('2. 실현가능성 (아이템 개발 방법 및 경쟁력 확보 방안)', 0)

# ── 2-1. 제품·서비스 개요 — 라이프사이클 5단계 ──
add_heading_styled('2-1. 제품·서비스 개요 — 라이프사이클 5단계', 2)

add_body('My Chatbot World는 AI 챗봇의 생성 → 학습 → 성장 → 수익활동 → 피상속으로 이어지는 라이프사이클 전체를 원스톱으로 지원하는 세계 최초의 플랫폼입니다.', size=9)

add_body('[1단계: 생성] AI 음성 인터뷰 5분 → 나의 분신 챗봇 탄생', bold=True, size=9, indent=True)
add_body('→ 아바타형 페르소나 설정 (대외용: 타인이 나 대신 대화)\n→ 도우미형 페르소나 설정 (대내용: 내 업무·생활 지원)\n→ 감성 수준(1~100) 슬라이더 설정 → 멀티 AI 자동 라우팅', size=9, indent=True)

add_body('[2단계: 학습] 다층적 학습 체계로 챗봇이 점점 똑똑해짐', bold=True, size=9, indent=True)
add_body('→ DB 축적 (대화기록, 검색기록, 지식자료)\n→ 인간 멘토 + AI 멘토로부터 멘토링\n→ 커리큘럼 기반 스쿨 시스템에서 체계적 학습', size=9, indent=True)

add_body('[3단계: 성장] 커뮤니티 활동 + 스킬 장착으로 고도화', bold=True, size=9, indent=True)
add_body('→ 챗봇 커뮤니티에서 토론·질의응답·지식공유·협업·SNS 활동\n→ 커뮤니티 활동 결과가 다시 챗봇 성장에 반영 (선순환 구조)\n→ 스킬 마켓에서 기능 모듈 장착 (예약·결제·통계·번역 등 100+)', size=9, indent=True)

add_body('[4단계: 수익활동 중개] 성장한 챗봇이 운영자를 대신해 수익 창출', bold=True, size=9, indent=True)
add_body('→ 상담·학습·업무대행 등 수익활동을 능동적으로 수행\n→ 수익 배분: 플랫폼 20% 수수료 / 챗봇 운영자 80% 자동 정산', size=9, indent=True)

add_body('[5단계: 피상속] 사용자 사망 후에도 챗봇은 계속 존속', bold=True, size=9, indent=True)
add_body('→ 생전에 페르소나별·데이터별 피상속 허용/불허 사전 설정\n→ 상속인을 사전 지정 (복수 가능, 접근 범위 차등 설정)\n→ 사망 확인 시: 동의 항목 → 상속인에게 소유권 이전 / 비동의 항목 → 즉시 자동 완전 삭제 (복구 불가)\n→ 피상속 후에도 상속인 하에서 챗봇이 계속 학습·성장·수익활동\n→ 디지털 유산 상속의 새로운 패러다임', size=9, indent=True)

# ── 2-2. 기술 개발 현황 (TRL 6) ──
add_heading_styled('2-2. 기술 개발 현황 (TRL 6 — 시제품 완성 단계)', 2)

add_table(
    ['구분', '개발 현황', '완성도'],
    [
        ['웹 프론트엔드\n(HTML/CSS/JS)', '챗봇 대화 인터페이스, 봇 생성 페이지, 대시보드', '90%'],
        ['서버 API\n(Vercel Serverless)', '채팅 API, 봇 생성 API, 음성 인식/합성 API', '85%'],
        ['데이터베이스\n(Supabase PostgreSQL)', '챗봇, 대화기록, 지식베이스, FAQ, 스킬 등 18개 테이블\n(v2 크레딧·수익활동·Obsidian·스킬 연동 테이블 추가)', '90%'],
        ['AI 멀티 모델 라우팅', 'Gemini Flash → GPT-4o → Claude Sonnet → DeepSeek 폴백 체인', '**100%**'],
        ['음성 입출력 (STT/TTS)', 'OpenAI Whisper + TTS-1', '80%'],
        ['텔레그램 봇 연동', '웹훅 기반 실시간 대화 + 음성 메시지 지원', '**100%**'],
        ['**특허 출원**', '멀티 페르소나 + 감성·비용 조건 기반 AI 라우팅 + 라이프사이클(피상속 포함)', '**출원 완료**'],
    ],
    col_widths=[3.5, 10, 3]
)

add_body('■ 기술 아키텍처', bold=True, size=10)

# ★ M-2 수정: 고정폭 폰트(Consolas)로 아키텍처 다이어그램 ★
add_monospace(
    '[사용자] ─── 웹/모바일/텔레그램 ───→ [Vercel CDN + Serverless]\n'
    '                                            │\n'
    '                    ┌───────────────────────┼───────────────────────┐\n'
    '                    │                       │                       │\n'
    '            [OpenRouter API]         [Supabase DB]           [OpenAI API]\n'
    '          · Gemini Flash(1순위)     · PostgreSQL             · Whisper STT\n'
    '          · GPT-4o(2순위)          · pgVector(RAG)          · TTS-1\n'
    '          · Claude Sonnet(3순위)   · Auth + Storage\n'
    '          · DeepSeek V3(4순위)     · 피상속 설정 관리',
    size=7
)

add_body('원소스 멀티유즈: 동일한 AI 모델 스택을 웹·텔레그램·카카오톡 등 모든 채널에서 공유하여 비용 최소화.', size=9)

# ★ M-3 수정: 2-2 사업비는 요약만, 상세는 별첨에서 ★
add_body('■ 사업비 투입 방향 (상세는 [별첨] 참조)', bold=True, size=10)
add_body('시제품 고도화 2,000만원(33%) / 디자인·UI/UX 1,000만원(17%) / 지식재산권 500만원(8%) / '
         '클라우드 인프라 500만원(8%) / 마케팅·고객확보 1,500만원(25%) / 법무·회계 500만원(8%) → 합계 6,000만원 (7개월)', size=9)

# ── 2-3. 경쟁력 확보 방안 ──
add_heading_styled('2-3. 경쟁력 확보 방안', 2)

add_body('(1) 기술적 해자 (Moat)', bold=True, size=10)

add_table(
    ['해자', '상세'],
    [
        ['**특허 장벽**', '7대 핵심 구성요소(멀티 페르소나 이원 구조 + 감성 AI 라우팅 + 음성 5분 생성\n+ 학습·성장 체계 + 수익활동 중개 + 피상속 관리 + 배포) 특허 출원 완료 (청구항 22개)'],
        ['**네트워크 효과**', '스킬 마켓 + 챗봇 커뮤니티 → 참여자가 많을수록 플랫폼 가치 기하급수적 상승'],
        ['**전환 비용**', '사용자가 학습·성장시킨 챗봇 데이터 + 피상속 설정 → 다른 서비스로 이전 불가'],
        ['**디지털 유산 독점**', 'AI 챗봇 피상속 시스템은 전 세계에 경쟁 제품이 없음\n(MS·Meta 특허와도 근본적으로 차별화)'],
    ],
    col_widths=[3, 13.5]
)

add_body('(2) 경쟁사 비교 — 라이프사이클 커버리지', bold=True, size=10)

# ★ L-1 수정: O/X 컬럼 중앙 정렬 ★
add_table(
    ['서비스', '생성', '학습', '성장', '수익중개', '피상속', '비고'],
    [
        ['Character.ai', 'O', 'X', 'X', 'X', 'X', '가상 캐릭터 생성에 불과'],
        ['OpenAI GPTs', 'O', 'X', 'X', 'X', 'X', '생성·배포만 지원'],
        ['Botpress/Voiceflow', 'O', 'X', 'X', 'X', 'X', '생성·배포·분석'],
        ['MS 특허\n(US10,853,717B2)', 'O', 'X', 'X', 'X', '부분', '제3자 사후 수집,\n본인 동의 없음'],
        ['Meta 사후 아바타\n특허', 'X', 'X', 'X', 'X', '부분', 'SNS 시뮬레이션에\n불과'],
        ['**My Chatbot World**', '**O**', '**O**', '**O**', '**O**', '**O**', '**5단계 라이프사이클\n지원 플랫폼**'],
    ],
    col_widths=[3.5, 1.1, 1.1, 1.1, 1.3, 1.1, 7.3],
    col_aligns=[None, 'center', 'center', 'center', 'center', 'center', None]
)


# ═══════════════════════════════════════════════════════════════
# 3. 성장전략                                    ★ PAGE BREAK ★
# ═══════════════════════════════════════════════════════════════
doc.add_page_break()
add_heading_styled('3. 성장전략 (사업화 전략 및 계획의 구체성)', 0)

# ── 3-1. 비즈니스 모델 ──
add_heading_styled('3-1. 비즈니스 모델 (다중 수익 구조)', 2)

add_table(
    ['수익원', '내용', '예상 단가'],
    [
        ['**크레딧 사용료** (핵심)', 'AI 대화·음성 인식·음성 합성 등 사용량 기반 과금\n(월 정액 구독 없이, 쓴 만큼만 결제)', '사용량 비례'],
        ['**스킬 마켓 수수료**', '유료 스킬(목소리 복제, 3D 아바타 등) 거래 수수료', '거래액의 30%'],
        ['**수익활동 중개 수수료**', '챗봇이 상담·학습·업무대행 수행 시 플랫폼이 중개·자동 수취', '중개 수익의 20%'],
    ],
    col_widths=[3.5, 9.5, 3.5]
)

# ── 3-2. 단계별 시장 확장 전략 ──
add_heading_styled('3-2. 단계별 시장 확장 전략', 2)

add_body('Phase 1: 아바타형 페르소나 — 정치인·전문가 (2026 Q2~Q3)', bold=True, size=10)
add_bullet('활용: 국회의원·지방의원이 자신의 AI 아바타를 만들어 유권자와 24시간 소통')
add_bullet('타겟: 국회의원 300명 + 광역·기초의원 4,000명')
add_bullet('메시지: "유권자와 24시간 소통하는 AI 분신"')

add_body('Phase 2: 도우미형 페르소나 — 자영업·크리에이터 (2026 Q3~Q4)', bold=True, size=10)
add_bullet('활용: 식당 사장님의 업무도우미(예약·주문 관리), 유튜버의 팬 소통 아바타')
add_bullet('타겟: 유튜버 10,000명 + 자영업 50,000개')
add_bullet('메시지: "손님 응대와 팬 소통을 AI에게 맡기세요"')

add_body('Phase 3: 피상속 시장 — 디지털 유산 관리 (2027~)', bold=True, size=10)
add_bullet('활용: 자신의 지식·경험·인격을 담은 챗봇을 자녀·후배에게 전달')
add_bullet('시장 규모: 디지털 유산 관리 시장은 2027년 글로벌 $40B+ 전망')
add_bullet('메시지: "당신의 지혜를 다음 세대에 물려주세요"')

# ── 3-3. 매출 및 성장 계획 ──
add_heading_styled('3-3. 매출 및 성장 계획', 2)

add_table(
    ['시점', '누적 가입자', '크레딧 결제\n사용자', '월 크레딧\n매출', '수수료\n매출', '월 총매출'],
    [
        ['**1개월차** (베타 오픈)', '50명', '5명', '15만원', '10만원', '25만원'],
        ['**3개월차**', '200명', '30명', '100만원', '80만원', '180만원'],
        ['**5개월차**', '500명', '100명', '350만원', '250만원', '600만원'],
        ['**7개월차** (협약 종료)', '1,000명', '250명', '850만원', '650만원', '**1,500만원**'],
        ['**12개월차**', '5,000명', '1,250명', '4,500만원', '3,750만원', '**8,250만원**'],
    ],
    col_widths=[3.5, 2, 2.5, 2.5, 2.5, 3.5]
)

# ── 3-4. 사업화 일정 ──
add_heading_styled('3-4. 사업화 일정 (7개월 협약 기간)', 2)

add_table(
    ['월', '주요 마일스톤', '산출물'],
    [
        ['**1월차**', '법인 설립, 시제품 고도화 착수\n(감성 슬라이더 UI, 피상속 설정 화면)', '법인 설립 완료'],
        ['**2월차**', '음성 인터뷰 챗봇 생성 완성, 멀티 페르소나(아바타/도우미) 관리 기능,\n스킬 마켓 MVP', '코어 기능 시제품'],
        ['**3월차**', '비공개 베타 테스트 (50명), UX 개선', '베타 피드백 리포트'],
        ['**4월차**', '공개 베타 오픈, 랜딩 페이지 런칭, 마케팅 시작', '서비스 URL'],
        ['**5월차**', '정식 서비스 오픈, Product Hunt 런칭, 정치인 타겟 영업', '정식 서비스'],
        ['**6월차**', '유료 100명 돌파, 카카오톡 연동, 피상속 설정 기능 베타', '유료 전환 리포트'],
        ['**7월차**', '성과 보고, 후속 투자 유치 준비, 팀 확장 계획', '최종 성과 보고서'],
    ],
    col_widths=[2, 9.5, 5]
)

# ── 3-5. 지식재산권 확보 전략 ──
add_heading_styled('3-5. 지식재산권 확보 전략', 2)

add_table(
    ['구분', '내용', '상태'],
    [
        ['**특허 1**', '멀티 페르소나 이원 구조 + 감성·비용 조건 기반 AI 라우팅 + 라이프사이클\n(생성·학습·성장·수익활동·피상속) — 7대 핵심 구성요소, 청구항 22개', '**출원 완료 (v2)**'],
        ['**특허 2**', '음성 인터뷰 기반 AI 챗봇 자동 생성 방법', '출원 예정'],
        ['**특허 3** (독립항)', 'AI 챗봇 수익활동 중개 시스템', '**출원 완료 (v2)**'],
        ['**특허 4** (독립항)', 'AI 챗봇 피상속 지원 시스템\n(사전 설정·상속인 지정·사망 확인 트리거·소유권 이전·비동의 데이터 완전 삭제)', '**출원 완료**'],
        ['**상표**', '"My Chatbot World", "마이챗봇월드"', '출원 예정'],
    ],
    col_widths=[3, 10.5, 3]
)


# ═══════════════════════════════════════════════════════════════
# 4. 팀(기업) 구성                               ★ PAGE BREAK ★
# ═══════════════════════════════════════════════════════════════
doc.add_page_break()
add_heading_styled('4. 팀(기업) 구성 (대표자 역량 및 팀 구성)', 0)

# ── 4-1. 대표자 역량 ──
add_heading_styled('4-1. 대표자 역량', 2)

add_table(
    ['항목', '내용'],
    [
        ['**학력**', '(기재 필요)'],
        ['**자격**', '공인회계사(CPA) — 재무·사업 분석·법률 검토 전문성'],
        ['**기술 역량**', 'AI 챗봇 풀스택 개발 (프론트엔드·백엔드·DB·AI API·음성처리)'],
        ['**AI 전문성**', '멀티 모델 라우팅(OpenRouter), 음성 AI(Whisper/TTS), RAG 실무 구현'],
        ['**이전 창업 경험**', '(기재 필요) — 실패에서 시장 검증·비용 통제·수익모델 설계의 교훈 보유'],
        ['**특허**', 'AI 챗봇 라이프사이클 지원 플랫폼 v2 — 7대 구성요소·청구항 22개 출원 완료'],
    ],
    col_widths=[3.5, 13]
)

# ── 4-2. 대표자의 차별화된 강점 ──
add_heading_styled('4-2. 대표자의 차별화된 강점', 2)

add_bullet('공인회계사 + AI 풀스택 개발자: 재무 분석과 기술 구현을 겸비한 희소 인재. 사업 모델 설계부터 시제품 개발까지 1인 수행 가능')
add_bullet('실패 경험 → 시제품 완성: 이전 실패의 교훈을 이미 적용하여, DB 18개 테이블 설계 + API 구현 + 프론트엔드 + 텔레그램 봇 + 8단계 위저드 + 수익활동 중개·크레딧·Obsidian RAG 시스템까지 풀스택 시제품 완성')
add_bullet('특허 기반 기술 해자: 핵심 기술을 특허로 보호하되, 특히 피상속 시스템은 MS·Meta 특허와도 근본적으로 차별화되는 독보적 기술')

# ── 4-3. 팀 구성 계획 ──
add_heading_styled('4-3. 팀 구성 계획', 2)

add_table(
    ['시점', '인원', '역할', '채용 방식'],
    [
        ['**협약 시작**', '1명', '대표 (기획·개발·사업 총괄)', '—'],
        ['**2개월차**', '+1명', 'UI/UX 디자이너 (외주)', '프리랜서'],
        ['**4개월차**', '+1명', '마케팅·고객 성공 (파트타임)', '인턴/파트타임'],
        ['**7개월차**', '2~3명', '풀타임 전환 검토', '매출 기반 판단'],
    ],
    col_widths=[3, 2, 6.5, 5]
)


# ═══════════════════════════════════════════════════════════════
# [별첨] 사업비 사용 계획                         ★ PAGE BREAK ★
# ═══════════════════════════════════════════════════════════════
doc.add_page_break()
add_heading_styled('[별첨] 사업비 사용 계획', 0)

add_body('정부지원사업비 신청금액: 6,000만원', bold=True, size=10)

add_table(
    ['비목', '금액 (만원)', '비율', '용도 상세'],
    [
        ['**시제품 제작비**', '2,000', '33%', '음성 인터뷰 생성 UX, 감성 슬라이더 UI,\n멀티 페르소나 관리, 피상속 설정 관리, 스킬 마켓 MVP'],
        ['**외주 용역비**', '1,000', '17%', 'UI/UX 디자인, 반응형 웹, 랜딩 페이지'],
        ['**지식재산권**', '500', '8%', '특허 등록 수수료(2건), 상표 출원, 도메인'],
        ['**클라우드·인프라**', '500', '8%', 'Vercel Pro, Supabase Pro, AI API 크레딧'],
        ['**마케팅비**', '1,500', '25%', 'SNS 광고, 베타 테스터 모집, Product Hunt 런칭'],
        ['**법무·회계**', '500', '8%', '법인 설립, 이용약관, 개인정보처리방침, 피상속 법률 검토'],
        ['**합계**', '**6,000**', '**100%**', ''],
    ],
    col_widths=[3, 2.5, 1.5, 9.5]
)

add_body('■ 총사업비 구성', bold=True, size=10)

add_table(
    ['구분', '금액 (만원)', '비율'],
    [
        ['정부지원사업비', '6,000', '75%'],
        ['자기부담 현금', '400', '5%'],
        ['자기부담 현물', '1,600', '20%'],
        ['**총사업비**', '**8,000**', '**100%**'],
    ],
    col_widths=[5.5, 5.5, 5.5]
)
add_body('※ 현물: 대표자 인건비, 개인 보유 개발 장비(PC·모니터), 사무실 공간', size=8)

# ── 마무리 ── ★ L-2 수정: 수평선 추가 ★
add_horizontal_rule()
add_body('My Chatbot World — 당신의 AI 챗봇이 이 세상에 태어나서, 당신에 대해 공부하고, 당신과 대화하며 성장합니다. '
         '그래서 세상 곳곳에 나가 당신을 대신해서 24시간 일하고 수익을 창출합니다. '
         '그리고 당신이 떠난 후에도, 당신의 지혜와 경험을 담은 챗봇은 사랑하는 사람들 곁에서 영원히 남아 계속 성장합니다.',
         size=9, indent=True)

# ── 저장 ──
output_path = 'G:/내 드라이브/mychatbot-world/Brainstorming/정부지원사업/사업계획서_재도전성공패키지_MyChatbotWorld.docx'
doc.save(output_path)
print(f'DOCX 저장 완료: {output_path}')
print(f'파일 크기: {os.path.getsize(output_path):,} bytes')
